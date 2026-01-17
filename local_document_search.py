import json
import faiss
import numpy as np
import os
import re
from typing import List, Dict, Tuple, Optional
from pathlib import Path
from sentence_transformers import SentenceTransformer
from groq import Groq

from dotenv import load_dotenv

load_dotenv()


try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. .docx files will be skipped.")


class LocalDocumentSemanticSearch:
    def __init__(self, 
                 embedding_model: str = "all-MiniLM-L6-v2", 
                 llm_model: str = "openai/gpt-oss-20b",
                 chunk_size: int = 1000, 
                 chunk_overlap: int = 200):
        """
        Initialize the local document semantic search system with local embeddings and Groq LLM
        
        Args:
            embedding_model: SentenceTransformer model for embeddings
            llm_model: Groq model for chunk selection
            chunk_size: Size of text chunks for processing
            chunk_overlap: Overlap between consecutive chunks
        """
        # Initialize local embedding model
        print(f"Loading embedding model: {embedding_model}")
        self.embedding_model = SentenceTransformer(embedding_model)
        
        # Initialize Groq client
        self.groq_client = Groq()
        self.llm_model = llm_model
        
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.documents = []
        self.document_chunks = []
        self.index = None
        self.dimension = None
        
        # Create embeddings directory
        self.embeddings_dir = "agno_document_qa/embeddings"
        os.makedirs(self.embeddings_dir, exist_ok=True)
        
        # File paths for saving/loading
        self.embeddings_file = os.path.join(self.embeddings_dir, "local_document_embeddings.npy")
        self.index_file = os.path.join(self.embeddings_dir, "local_faiss_index.bin")
        self.chunks_file = os.path.join(self.embeddings_dir, "local_document_chunks.json")
        
        # Load documents from the documents folder if it exists
        self._load_documents()
        
        # Build or load the search index
        if self.documents:
            self._build_or_load_index()
    
    def _load_documents(self):
        """Load documents from various file formats"""
        documents_dir = "agno_document_qa/documents"
        if not os.path.exists(documents_dir):
            os.makedirs(documents_dir, exist_ok=True)
            print(f"Created documents directory: {documents_dir}")
            print("Please add your documents (.txt, .md, .json, .docx, .py, .js, .html, .css) to this folder")
            return
        
        supported_extensions = ['.txt', '.md', '.json', '.py', '.js', '.html', '.css', '.docx']
        
        for file_path in Path(documents_dir).rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    if file_path.suffix.lower() == '.docx':
                        content = self._extract_docx_content(file_path)
                    else:
                        with open(file_path, 'r', encoding='utf-8') as f:
                            content = f.read()
                    
                    self.documents.append({
                        'filename': file_path.name,
                        'filepath': str(file_path),
                        'content': content,
                        'size': len(content)
                    })
                    print(f"Loaded document: {file_path.name}")
                except Exception as e:
                    print(f"Error loading {file_path}: {e}")
        
        print(f"Total documents loaded: {len(self.documents)}")
    
    def _extract_docx_content(self, file_path: Path) -> str:
        """Extract text content from .docx files"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx package is required to read .docx files. Install with: pip install python-docx")
        
        try:
            doc = Document(file_path)
            content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content.append(" | ".join(row_text))
            
            return "\n\n".join(content)
            
        except Exception as e:
            raise Exception(f"Error extracting content from {file_path.name}: {str(e)}")
    
    def _chunk_text(self, text: str, metadata: Dict) -> List[Dict]:
        """Split text into overlapping chunks"""
        chunks = []
        
        # Simple sentence-aware chunking
        sentences = re.split(r'[.!?]+', text)
        current_chunk = ""
        current_size = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            sentence_size = len(sentence)
            
            # If adding this sentence exceeds chunk size, save current chunk
            if current_size + sentence_size > self.chunk_size and current_chunk:
                chunks.append({
                    'text': current_chunk.strip(),
                    'metadata': metadata.copy(),
                    'chunk_id': len(chunks)
                })
                
                # Start new chunk with overlap
                overlap_text = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else current_chunk
                current_chunk = overlap_text + " " + sentence
                current_size = len(current_chunk)
            else:
                current_chunk += " " + sentence if current_chunk else sentence
                current_size += sentence_size
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append({
                'text': current_chunk.strip(),
                'metadata': metadata.copy(),
                'chunk_id': len(chunks)
            })
        
        return chunks
    
    def _process_documents(self):
        """Process all documents into chunks"""
        self.document_chunks = []
        
        for doc in self.documents:
            metadata = {
                'filename': doc['filename'],
                'filepath': doc['filepath'],
                'document_size': doc['size']
            }
            
            chunks = self._chunk_text(doc['content'], metadata)
            self.document_chunks.extend(chunks)
        
        print(f"Created {len(self.document_chunks)} chunks from {len(self.documents)} documents")
    
    def _get_local_embeddings(self, texts: List[str]) -> np.ndarray:
        """Get embeddings using local SentenceTransformer model"""
        try:
            print(f"Generating embeddings for {len(texts)} texts using local model...")
            embeddings = self.embedding_model.encode(texts, show_progress_bar=True)
            return embeddings.astype('float32')
        except Exception as e:
            print(f"Error getting local embeddings: {e}")
            # Fallback to random embeddings for testing
            return np.random.rand(len(texts), 384).astype('float32')
    
    def _build_or_load_index(self):
        """Build or load FAISS index with document embeddings"""
        # Check if saved embeddings exist
        if (os.path.exists(self.embeddings_file) and 
            os.path.exists(self.index_file) and 
            os.path.exists(self.chunks_file)):
            
            try:
                self._load_index()
                print(f"Loaded existing local embeddings for {len(self.document_chunks)} chunks")
                return
            except Exception as e:
                print(f"Error loading existing embeddings: {e}")
                print("Rebuilding embeddings...")
        
        # Build new embeddings
        self._build_index()
    
    def _build_index(self):
        """Build FAISS index with document embeddings using local models"""
        print("Building document embeddings and FAISS index with local models...")
        
        if not self.documents:
            print("No documents found! Please add documents to the documents folder.")
            return
        
        # Process documents into chunks
        self._process_documents()
        
        if not self.document_chunks:
            print("No chunks created from documents!")
            return
        
        # Extract text from chunks for embedding
        chunk_texts = [chunk['text'] for chunk in self.document_chunks]
        
        print(f"Generating embeddings for {len(chunk_texts)} document chunks using local model...")
        
        # Generate embeddings using local model
        embeddings = self._get_local_embeddings(chunk_texts)
        self.dimension = embeddings.shape[1]
        
        # Create FAISS index (using L2 distance)
        self.index = faiss.IndexFlatL2(self.dimension)
        
        # Add embeddings to index
        self.index.add(embeddings)
        
        # Save embeddings and index to disk
        self._save_index(embeddings)
        
        print(f"Successfully indexed {len(self.document_chunks)} document chunks")
        print(f"Embedding dimension: {self.dimension}")
        print(f"Saved embeddings to {self.embeddings_dir}/")
    
    def _save_index(self, embeddings: np.ndarray):
        """Save embeddings and FAISS index to disk"""
        try:
            # Save embeddings as numpy array
            np.save(self.embeddings_file, embeddings)
            
            # Save FAISS index
            faiss.write_index(self.index, self.index_file)
            
            # Save document chunks for reference
            with open(self.chunks_file, 'w', encoding='utf-8') as f:
                json.dump(self.document_chunks, f, indent=2, ensure_ascii=False)
            
            print(f"Saved embeddings to: {self.embeddings_file}")
            print(f"Saved FAISS index to: {self.index_file}")
            print(f"Saved document chunks to: {self.chunks_file}")
            
        except Exception as e:
            print(f"Error saving embeddings: {e}")
    
    def _load_index(self):
        """Load embeddings and FAISS index from disk"""
        # Load embeddings
        embeddings = np.load(self.embeddings_file)
        self.dimension = embeddings.shape[1]
        
        # Load FAISS index
        self.index = faiss.read_index(self.index_file)
        
        # Load document chunks
        with open(self.chunks_file, 'r', encoding='utf-8') as f:
            self.document_chunks = json.load(f)
        
        print(f"Loaded {len(self.document_chunks)} chunks from saved local index")
    
    def search_documents(self, query: str, top_k: int = 10) -> List[Dict]:
        """
        Search for relevant document chunks based on user query using local models
        
        Args:
            query: User's question or search query
            top_k: Number of top chunks to return
            
        Returns:
            List of relevant document chunks with similarity scores
        """
        if not self.index:
            raise ValueError("Search index not built. No documents available.")
        
        # Generate embedding for the query using local model
        query_embedding = self._get_local_embeddings([query])
        
        # Search in FAISS index
        distances, indices = self.index.search(query_embedding, top_k)
        
        # Retrieve matching chunks with similarity scores
        results = []
        for i, (distance, idx) in enumerate(zip(distances[0], indices[0])):
            if idx < len(self.document_chunks):  # Ensure valid index
                chunk = self.document_chunks[idx].copy()
                # Convert L2 distance to similarity score (lower distance = higher similarity)
                similarity_score = 1 / (1 + distance)
                chunk['similarity_score'] = float(similarity_score)
                chunk['rank'] = i + 1
                results.append(chunk)
        
        return results
    
    def _select_best_chunks_with_groq_llm(self, query: str, search_results: List[Dict]) -> List[Dict]:
        """
        Use Groq LLM to choose the best document chunks from search results
        
        Args:
            query: User's original query
            search_results: List of document chunks from semantic search
            
        Returns:
            List of selected document chunks
        """
        if not search_results:
            return []
        
        # Format search results for LLM analysis
        results_text = ""
        for i, chunk in enumerate(search_results, 1):
            results_text += f"{i}. From {chunk['metadata']['filename']}\n"
            results_text += f"   Content: {chunk['text'][:200]}...\n"
            results_text += f"   Similarity: {chunk['similarity_score']:.3f}\n\n"
        
        prompt = f"""Given the user query: "{query}"

And these document chunks:
{results_text}

Your task is to select the most relevant chunk numbers for answering the user's query.

**INSTRUCTIONS:**
1. Analyze which chunks contain information most relevant to the query
2. Select 2-3 chunks that best answer the question
3. Consider both content relevance and similarity scores
4. Return only the numbers of selected chunks, comma-separated

**OUTPUT FORMAT:**
Return only the chunk numbers (e.g., "1,3,5" or "2" for single chunk)
If no chunks are relevant, return "None"

Selected chunks:"""

        try:
            response = self.groq_client.chat.completions.create(
                model=self.llm_model,
                messages=[
                    {"role": "user", "content": prompt}
                ],
                temperature=0.1,
                max_tokens=50
            )
            
            choice_text = response.choices[0].message.content.strip()
            print(f"Groq LLM chunk selection: {choice_text}")
            
            if choice_text.lower() == 'none':
                return []
            
            try:
                # Extract numbers from response
                import re
                numbers = re.findall(r'\d+', choice_text)
                choice_nums = [int(num) for num in numbers]
            except (ValueError, AttributeError):
                print(f"Warning: Could not parse '{choice_text}' as numbers. Using top 3 results.")
                return search_results[:3]
            
            # Return selected chunks
            valid_results = [search_results[num - 1] for num in choice_nums if 1 <= num <= len(search_results)]
            return valid_results if valid_results else search_results[:3]
            
        except Exception as e:
            print(f"Error in Groq LLM selection: {e}")
            # Fallback to top 3 results
            return search_results[:3]
    
    def get_context_for_query(self, query: str, max_chunks: int = 3) -> str:
        """
        Get relevant context from documents for answering a query using local embeddings and Groq LLM
        
        Args:
            query: User's question
            max_chunks: Maximum number of chunks to include in context
            
        Returns:
            Formatted context string for the LLM
        """
        try:
            # Search for relevant chunks (get more for LLM selection)
            top_10_results = self.search_documents(query, top_k=10)
            
            # Use Groq LLM to select the best chunks
            selected_chunks = self._select_best_chunks_with_groq_llm(query, top_10_results)
            
            if not selected_chunks:
                return "No relevant information found in the documents."
            
            # Format context
            context_parts = []
            for i, chunk in enumerate(selected_chunks[:max_chunks], 1):
                context_parts.append(f"--- Document Chunk {i} (from {chunk['metadata']['filename']}) ---")
                context_parts.append(chunk['text'])
                context_parts.append("")  # Empty line for separation
            
            return "\n".join(context_parts)
            
        except Exception as e:
            return f"Error retrieving context: {str(e)}"


# Global instance for use in tools
_local_search_engine = None

def initialize_local_search_engine():
    """Initialize the global local search engine instance"""
    global _local_search_engine
    if _local_search_engine is None:
        _local_search_engine = LocalDocumentSemanticSearch()
    return _local_search_engine

def search_local_documents(query: str) -> str:
    """
    Tool function for the Agno agent to search documents using local embeddings and Groq LLM
    
    Args:
        query: User's question or search query
        
    Returns:
        JSON string with relevant document context
    """
    try:
        search_engine = initialize_local_search_engine()
        
        if not search_engine.documents:
            return json.dumps({
                "error": "No documents available. Please add documents to the agno_document_qa/documents folder.",
                "context": "",
                "query": query
            })
        
        # Get relevant context using local embeddings and Groq LLM
        context = search_engine.get_context_for_query(query, max_chunks=3)
        
        # Also get the raw search results for metadata
        search_results = search_engine.search_documents(query, top_k=3)
        
        return json.dumps({
            "query": query,
            "context": context,
            "sources": [
                {
                    "filename": chunk['metadata']['filename'],
                    "similarity_score": chunk['similarity_score'],
                    "rank": chunk['rank']
                }
                for chunk in search_results
            ],
            "total_chunks_found": len(search_results),
            "model_info": {
                "embedding_model": "SentenceTransformer (local)",
                "llm_model": search_engine.llm_model + " (Groq)"
            }
        }, indent=2)
        
    except Exception as e:
        return json.dumps({
            "error": f"Failed to search documents with local embeddings and Groq LLM: {str(e)}",
            "context": "",
            "query": query
        })


# Test function
def test_local_document_search():
    """Test the local document search functionality"""
    search_engine = LocalDocumentSemanticSearch()
    
    if not search_engine.documents:
        print("No documents found for testing. Please add some documents to the documents folder.")
        return
    
    test_queries = [
        "What is the main topic of the documents?",
        "How does this system work?",
        "What are the key features?",
        "Can you explain the implementation?",
        "What technologies are used?"
    ]
    
    print("Testing Local Document Search:")
    print("=" * 50)
    
    for query in test_queries:
        print(f"\nQuery: {query}")
        results = search_engine.search_documents(query, top_k=2)
        
        for i, chunk in enumerate(results, 1):
            print(f"{i}. From {chunk['metadata']['filename']}")
            print(f"   Score: {chunk['similarity_score']:.3f}")
            print(f"   Preview: {chunk['text'][:100]}...")

if __name__ == "__main__":
    test_local_document_search()
